"""Session report builder: collect analysis results and export them as a Word document."""
import io

import pandas as pd
import streamlit as st


def init_report_state():
    if "report" not in st.session_state:
        st.session_state.report = []


def add_report_entry(title: str, blocks: list):
    """blocks: list of ("text" | "table", content) tuples, content is str or pd.DataFrame."""
    init_report_state()
    st.session_state.report.append({"title": title, "blocks": blocks})


def clear_report():
    init_report_state()
    st.session_state.report = []


def report_count() -> int:
    init_report_state()
    return len(st.session_state.report)


def _add_df_table(doc, df: pd.DataFrame):
    from docx.shared import Pt

    df = df.reset_index() if df.index.name or not isinstance(df.index, pd.RangeIndex) else df
    table = doc.add_table(rows=1, cols=len(df.columns))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        pass
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if pd.isna(val) else str(val)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)


def build_docx_bytes(report_title: str = "DASS — Analysis Report") -> bytes:
    from docx import Document

    init_report_state()
    doc = Document()
    doc.add_heading(report_title, level=1)

    for entry in st.session_state.report:
        doc.add_heading(entry["title"], level=2)
        for block_type, content in entry["blocks"]:
            if block_type == "text":
                doc.add_paragraph(content)
            elif block_type == "table" and isinstance(content, pd.DataFrame):
                _add_df_table(doc, content)
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

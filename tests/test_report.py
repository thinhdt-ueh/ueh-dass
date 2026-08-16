import io

import pandas as pd
import streamlit as st

from utils.report import add_report_entry, build_docx_bytes, clear_report, report_count


def test_report_count_and_clear():
    st.session_state.report = []
    assert report_count() == 0
    add_report_entry("Entry 1", [("text", "hello")])
    assert report_count() == 1
    clear_report()
    assert report_count() == 0


def test_build_docx_with_text_and_table():
    st.session_state.report = []
    df = pd.DataFrame({"A": [1, 2], "B": ["x", "y"]})
    add_report_entry("Test Entry", [("text", "some text"), ("table", df)])

    docx_bytes = build_docx_bytes()
    assert len(docx_bytes) > 500

    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    assert len(doc.tables) == 1
    assert doc.tables[0].rows[0].cells[0].text == "A"


def test_build_docx_with_image_block():
    st.session_state.report = []
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (4, 4), color="white").save(buf, format="PNG")
    add_report_entry("Image Entry", [("image", buf.getvalue())])

    docx_bytes = build_docx_bytes()
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    assert len(doc.inline_shapes) == 1
